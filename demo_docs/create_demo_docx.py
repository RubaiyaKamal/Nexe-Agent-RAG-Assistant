"""Generate demo_company_handbook.docx for RAG Assistant testing."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Nexe Corp — Employee Handbook 2024", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Welcome to Nexe Corp. This handbook outlines our company policies, culture, "
    "benefits, and procedures. Please read it carefully and keep it as a reference."
)

# ── Section 1 ──────────────────────────────────────────────────────────────
doc.add_heading("1. Company Overview", level=1)
doc.add_paragraph(
    "Nexe Corp was founded in 2018 with a mission to build intelligent software tools "
    "that help teams move faster and smarter. We are headquartered in San Francisco, CA, "
    "with remote offices in London, Tokyo, and Nairobi."
)
doc.add_paragraph(
    "Our core products include the Nexe Agent Platform, a suite of AI-powered automation "
    "tools, and the Nexe RAG Assistant, an enterprise knowledge-retrieval system."
)

# ── Section 2 ──────────────────────────────────────────────────────────────
doc.add_heading("2. Work Hours & Remote Policy", level=1)
doc.add_paragraph(
    "Nexe Corp operates on a flexible, results-oriented work schedule. There are no fixed "
    "core hours; however, employees are expected to be available for team syncs and "
    "cross-timezone collaboration windows:"
)
bullet_items = [
    "Americas overlap: 10:00 AM – 2:00 PM Pacific Time",
    "EMEA overlap: 9:00 AM – 1:00 PM GMT",
    "APAC overlap: 9:00 AM – 1:00 PM JST",
]
for item in bullet_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "All full-time employees may work remotely. Those based within 50 miles of a Nexe "
    "office are encouraged (but not required) to come in 2 days per week."
)

# ── Section 3 ──────────────────────────────────────────────────────────────
doc.add_heading("3. Compensation & Benefits", level=1)

doc.add_heading("3.1 Salary Review", level=2)
doc.add_paragraph(
    "Salaries are reviewed annually in January. Performance ratings from Q4 reviews "
    "directly inform merit increases. The company targets the 65th percentile of market "
    "compensation benchmarks for all roles."
)

doc.add_heading("3.2 Health Insurance", level=2)
doc.add_paragraph(
    "Nexe Corp covers 100% of employee premiums and 80% of dependent premiums for medical, "
    "dental, and vision insurance. Coverage begins on the first day of employment."
)

doc.add_heading("3.3 Equity", level=2)
doc.add_paragraph(
    "All full-time employees receive stock options with a 4-year vesting schedule and a "
    "1-year cliff. Option grants are reviewed and refreshed annually based on performance."
)

doc.add_heading("3.4 Other Benefits", level=2)
perks = [
    "$500/month home office stipend for remote employees",
    "$2,000/year learning & development budget",
    "Unlimited PTO (with a minimum 10-day usage requirement)",
    "16 weeks paid parental leave for all parents",
    "$100/month wellness reimbursement (gym, therapy, meditation apps)",
    "401(k) with 4% company match (US employees)",
]
for perk in perks:
    doc.add_paragraph(perk, style="List Bullet")

# ── Section 4 ──────────────────────────────────────────────────────────────
doc.add_heading("4. Code of Conduct", level=1)
doc.add_paragraph(
    "Nexe Corp is committed to creating an inclusive, respectful, and psychologically safe "
    "workplace. All employees, contractors, and partners are expected to:"
)
conduct_items = [
    "Treat colleagues with dignity and respect regardless of background, identity, or role.",
    "Communicate constructively and assume positive intent.",
    "Protect confidential company and customer information.",
    "Report conflicts of interest promptly to HR.",
    "Never engage in harassment, discrimination, or retaliation.",
]
for item in conduct_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "Violations of the code of conduct may result in disciplinary action up to and including "
    "termination. Employees are encouraged to report concerns via the anonymous HR hotline "
    "at ethics@nexecorp.example.com."
)

# ── Section 5 ──────────────────────────────────────────────────────────────
doc.add_heading("5. Performance Reviews", level=1)
doc.add_paragraph(
    "Performance is evaluated twice per year: a mid-year check-in in June and a full annual "
    "review in December. Reviews assess:"
)
review_items = [
    "Goal achievement (40%)",
    "Core competencies and skills (30%)",
    "Collaboration and culture contribution (20%)",
    "Growth and development (10%)",
]
for item in review_items:
    doc.add_paragraph(item, style="List Number")

doc.add_paragraph(
    "Employees rate themselves first; managers then provide independent ratings. Final "
    "calibration sessions ensure consistency across teams."
)

# ── Section 6 ──────────────────────────────────────────────────────────────
doc.add_heading("6. IT & Security Policies", level=1)
doc.add_paragraph(
    "All employees must complete annual security awareness training. Key policies include:"
)
sec_items = [
    "Use strong, unique passwords managed via a company-approved password manager.",
    "Enable MFA on all company accounts.",
    "Never share credentials or leave sessions unlocked on shared machines.",
    "Report phishing attempts or suspected breaches to security@nexecorp.example.com immediately.",
    "All company data must be stored in approved cloud services (Google Workspace, Notion, GitHub).",
]
for item in sec_items:
    doc.add_paragraph(item, style="List Bullet")

# ── Section 7 ──────────────────────────────────────────────────────────────
doc.add_heading("7. Onboarding Checklist", level=1)
doc.add_paragraph("New employees should complete the following in their first two weeks:")
onboarding = [
    "Sign and return offer letter, NDA, and equity paperwork.",
    "Set up company laptop using the IT setup guide (IT Wiki → New Employee Setup).",
    "Complete security awareness training module.",
    "Schedule 1:1 introductions with your manager and immediate team.",
    "Join relevant Slack channels: #general, #your-team, #announcements.",
    "Review and acknowledge this handbook.",
    "Set up your 30-60-90 day goals with your manager.",
]
for item in onboarding:
    doc.add_paragraph(item, style="List Number")

# ── Footer note ────────────────────────────────────────────────────────────
doc.add_paragraph("")
p = doc.add_paragraph(
    "This handbook is reviewed and updated annually. Last updated: January 2024. "
    "Questions? Contact hr@nexecorp.example.com."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.runs[0]
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.size = Pt(9)

out_path = "demo_company_handbook.docx"
doc.save(out_path)
print(f"Created: {out_path}")
